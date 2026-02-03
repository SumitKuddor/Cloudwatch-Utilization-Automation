import boto3
import json
import os
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

from openpyxl import Workbook
from openpyxl.styles import Font

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# =====================================================
# CONFIGURATION
# =====================================================
REGION = "ap-south-1"
IMG_DIR = "images"

IST = ZoneInfo("Asia/Kolkata")
END = datetime.now(tz=IST)
START = END - timedelta(days=1)

TODAY = END.strftime("%d-%m-%Y")

EXCEL_FILE = f"reports/excel/Daily_Utilization_{TODAY}.xlsx"
WORD_FILE  = f"reports/word/Daily_Utilization_{TODAY}.docx"

# =====================================================
# DIRECTORIES
# =====================================================
os.makedirs(IMG_DIR, exist_ok=True)
os.makedirs("reports/excel", exist_ok=True)
os.makedirs("reports/word", exist_ok=True)

# =====================================================
# AWS CLIENTS
# =====================================================
ec2 = boto3.client("ec2", region_name=REGION)
rds = boto3.client("rds", region_name=REGION)
cw  = boto3.client("cloudwatch", region_name=REGION)

# =====================================================
# METRIC AVERAGE FETCH
# =====================================================
def avg_metric(namespace, metric, dimensions):
    try:
        res = cw.get_metric_statistics(
            Namespace=namespace,
            MetricName=metric,
            Dimensions=dimensions,
            StartTime=START,
            EndTime=END,
            Period=300,
            Statistics=["Average"]
        )
        if not res["Datapoints"]:
            return "NA"
        return round(
            sum(d["Average"] for d in res["Datapoints"]) / len(res["Datapoints"]),
            2
        )
    except:
        return "NA"

# =====================================================
# CLOUDWATCH GRAPH IMAGE (HIGH RES)
# =====================================================
def save_graph(widget, outfile):
    try:
        widget["width"] = 1200
        widget["height"] = 350

        img = cw.get_metric_widget_image(
            MetricWidget=json.dumps(widget)
        )["MetricWidgetImage"]

        with open(outfile, "wb") as f:
            f.write(img)

        return outfile
    except:
        return None

# =====================================================
# EC2 PLATFORM DETECTION
# =====================================================
def detect_platform(instance):
    return "windows" if instance.get("Platform") == "windows" else "linux"

# =====================================================
# LINUX DISK DIMENSIONS
# =====================================================
def get_linux_disk_dimensions(iid):
    try:
        metrics = cw.list_metrics(
            Namespace="CWAgent",
            MetricName="disk_used_percent",
            Dimensions=[{"Name": "InstanceId", "Value": iid}]
        )["Metrics"]

        return metrics[0]["Dimensions"] if metrics else None
    except:
        return None

# =====================================================
# EC2 MEMORY
# =====================================================
def get_ec2_memory(iid, platform):
    if platform == "windows":
        return avg_metric(
            "CWAgent",
            "Memory % Committed Bytes In Use",
            [{"Name": "InstanceId", "Value": iid}]
        )
    else:
        return avg_metric(
            "CWAgent",
            "mem_used_percent",
            [{"Name": "InstanceId", "Value": iid}]
        )

# =====================================================
# EC2 DISK
# =====================================================
def get_ec2_disk(iid, platform):
    if platform == "windows":
        free = avg_metric(
            "CWAgent",
            "LogicalDisk % Free Space",
            [{"Name": "InstanceId", "Value": iid}]
        )
        return "NA" if free == "NA" else round(100 - free, 2)
    else:
        dims = get_linux_disk_dimensions(iid)
        return avg_metric("CWAgent", "disk_used_percent", dims) if dims else "NA"

# =====================================================
# FETCH EC2
# =====================================================
def fetch_ec2():
    data = {}
    for page in ec2.get_paginator("describe_instances").paginate():
        for res in page["Reservations"]:
            for inst in res["Instances"]:
                iid = inst["InstanceId"]
                data[iid] = {
                    "name": next((t["Value"] for t in inst.get("Tags", []) if t["Key"] == "Name"), "NA"),
                    "state": inst["State"]["Name"].upper(),
                    "platform": detect_platform(inst),
                    "public_ip": inst.get("PublicIpAddress", "NA"),
                    "private_ip": inst.get("PrivateIpAddress", "NA")
                }
    return data

# =====================================================
# FETCH RDS
# =====================================================
def fetch_rds():
    return rds.describe_db_instances()["DBInstances"]

# =====================================================
# WORD BORDERED BOX
# =====================================================
def bordered_box(doc, text=None, img=None):
    table = doc.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl.tblPr.append(borders)

    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if img:
        p.add_run().add_picture(img, width=Inches(6), height=Inches(1.7))
    else:
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")

# =====================================================
# WORD REPORT
# =====================================================
def generate_word(ec2_data, rds_data):
    doc = Document()

    doc.add_heading("AWS Daily Utilization Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date : {TODAY}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---------------- EC2 SECTION ----------------
    for iid, i in ec2_data.items():
        p = doc.add_paragraph(i["name"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(15)

        doc.add_paragraph(iid).alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i["state"] != "RUNNING":
            bordered_box(doc, text="STOPPED")
            doc.add_page_break()
            continue

        doc.add_paragraph("1) CPU UTILIZATION").runs[0].bold = True
        bordered_box(doc, img=save_graph({
            "metrics": [["AWS/EC2", "CPUUtilization", "InstanceId", iid]],
            "region": REGION,
            "stat": "Average",
            "period": 300,
            "start": START.isoformat(),
            "end": END.isoformat()
        }, f"{IMG_DIR}/{iid}_cpu.png"))

        doc.add_paragraph("2) MEMORY UTILIZATION").runs[0].bold = True
        mem_metric = "Memory % Committed Bytes In Use" if i["platform"] == "windows" else "mem_used_percent"
        bordered_box(doc, img=save_graph({
            "metrics": [["CWAgent", mem_metric, "InstanceId", iid]],
            "region": REGION,
            "stat": "Average",
            "period": 300,
            "start": START.isoformat(),
            "end": END.isoformat()
        }, f"{IMG_DIR}/{iid}_mem.png"))

        doc.add_paragraph("3) DISK UTILIZATION").runs[0].bold = True
        if i["platform"] == "windows":
            disk_widget = {
                "metrics": [["CWAgent", "LogicalDisk % Free Space", "InstanceId", iid]]
            }
        else:
            dims = get_linux_disk_dimensions(iid)
            disk_widget = {
                "metrics": [["CWAgent", "disk_used_percent"] + sum([[d["Name"], d["Value"]] for d in dims], [])]
            } if dims else None

        bordered_box(doc, img=save_graph({
            **disk_widget,
            "region": REGION,
            "stat": "Average",
            "period": 300,
            "start": START.isoformat(),
            "end": END.isoformat()
        }, f"{IMG_DIR}/{iid}_disk.png") if disk_widget else None)

        doc.add_page_break()

    # ---------------- RDS SECTION ----------------
    doc.add_heading("AWS RDS Utilization Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date : {TODAY}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    for db in rds_data:
        dbid = db["DBInstanceIdentifier"]
        status = db["DBInstanceStatus"].upper()

        p = doc.add_paragraph(dbid)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        if status != "AVAILABLE":
            bordered_box(doc, text="STOPPED")
            doc.add_page_break()
            continue

        for title, metric in [
            ("CPU UTILIZATION", "CPUUtilization"),
            ("FREEABLE MEMORY", "FreeableMemory"),
            ("DB CONNECTIONS", "DatabaseConnections"),
            ("READ IOPS", "ReadIOPS"),
            ("WRITE IOPS", "WriteIOPS")
        ]:
            doc.add_paragraph(title).runs[0].bold = True
            bordered_box(doc, img=save_graph({
                "metrics": [["AWS/RDS", metric, "DBInstanceIdentifier", dbid]],
                "region": REGION,
                "stat": "Average",
                "period": 300,
                "start": START.isoformat(),
                "end": END.isoformat()
            }, f"{IMG_DIR}/{dbid}_{metric}.png"))

        doc.add_page_break()

    doc.save(WORD_FILE)

# =====================================================
# MAIN
# =====================================================
def main():
    ec2_data = fetch_ec2()
    rds_data = fetch_rds()

    generate_word(ec2_data, rds_data)

    for f in os.listdir(IMG_DIR):
        os.remove(os.path.join(IMG_DIR, f))

    print("✅ FINAL EC2 + RDS WORD REPORT GENERATED (CLEAR GRAPHS)")

if __name__ == "__main__":
    main()

